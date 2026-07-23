# Verify every code snippet shown in book/ch17.js.
# Same self-checking idea as ch2_verify.py: each snippet's expected output (as
# embedded in the chapter file) is asserted here, so any drift between the
# chapter text and real execution fails the run.
#
# Stdlib checks always run. The pypdf / python-docx checks need those packages
# (pip install pypdf python-docx); when they are missing, that section prints a
# skip line and the run still exits 0. No network. All document files are
# rebuilt from scratch in a temp directory, the same way they were built when
# the chapter's outputs were captured.
import io, contextlib, os, sys, tempfile, struct, zlib

fails = 0
total = 0
def check(label, got, want):
    global fails, total
    total += 1
    ok = got == want
    print(f"{'ok  ' if ok else 'FAIL'} {label}")
    if not ok:
        fails += 1
        print(f"     got:  {got!r}")
        print(f"     want: {want!r}")

def run(code):
    # Execute a chapter code block verbatim, capturing stdout.
    buf = io.StringIO()
    ns = {}
    with contextlib.redirect_stdout(buf):
        exec(code, ns)
    return buf.getvalue().rstrip("\n")

print("=== stdlib: graded exercises, reference solutions vs expected o ===")
check("ex1 Reassemble the runs",
    run("""runs = ['Rot', 'a posted', ' on the shed door.']
print(''.join(runs))
print(len(runs))"""),
    "Rota posted on the shed door.\n3")
check("ex2 Full text, minus the blanks",
    run("""paras = ['Minutes', '', 'All present.', '', 'Meeting closed.']
for p in paras:
    if p != '':
        print(p)"""),
    "Minutes\nAll present.\nMeeting closed.")
check("ex3 Word counts by page",
    run("""pages = ['Harbour Rowing Club minutes', 'Item one covers the boat shed roof', 'The regatta lands in late June']
total = 0
for page in pages:
    n = len(page.split())
    print(n)
    total += n
print(total)"""),
    "4\n7\n6\n17")
check("ex4 Build the table of contents",
    run("""toc = [(1, 'Getting the plot'), (2, 'Tools'), (2, 'Watering'), (1, 'First harvest')]
for level, title in toc:
    print('  ' * (level - 1) + title)"""),
    "Getting the plot\n  Tools\n  Watering\nFirst harvest")

print("=== stdlib: pure-Python claims from sections, questions, hints ===")
check("q8 claim: para.text is ''.join of the run pieces",
    ''.join(['Wat', 'er butts', ' refill on Fridays.']),
    'Water butts refill on Fridays.')
check("ex4 hint claim: '  ' * 2 is four spaces", '  ' * 2, '    ')
check("ex4 hint claim: '  ' * 0 is the empty string", '  ' * 0, '')
check("section 6 claim: newline-join of the six paragraph texts",
    '\n'.join(['The Plot Holder', 'Compost rota',
               'Turning duty passes to bed seven this month.',
               'Deliveries arrive Thursday at the north gate.',
               'Seed swap', 'Bring labelled envelopes to the shed.']),
    "The Plot Holder\nCompost rota\nTurning duty passes to bed seven this month.\n"
    "Deliveries arrive Thursday at the north gate.\nSeed swap\nBring labelled envelopes to the shed.")

# ---------------------------------------------------------------------------
# Guarded section: everything below needs pypdf and python-docx.
GUARDED_CHECKS = 47
try:
    import pypdf, docx  # noqa: F401 - probing availability of both
    HAVE_LIBS = True
except ImportError:
    HAVE_LIBS = False

def make_pdf(path, pages, font_size=18):
    # Handmade minimal PDF (correct xref) - same generator that built the
    # sample files whose outputs the chapter shows.
    n_pages = len(pages)
    page_ids = [3 + i for i in range(n_pages)]
    content_ids = [3 + n_pages + i for i in range(n_pages)]
    font_id = 3 + 2 * n_pages
    def esc(s):
        return s.replace('\\', r'\\').replace('(', r'\(').replace(')', r'\)')
    body = {}
    body[1] = "<< /Type /Catalog /Pages 2 0 R >>"
    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    body[2] = f"<< /Type /Pages /Kids [{kids}] /Count {n_pages} >>"
    for i, lines in enumerate(pages):
        body[page_ids[i]] = (
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
            f"/Contents {content_ids[i]} 0 R >>")
        parts = ["BT", f"/F1 {font_size} Tf", f"{font_size + 4} TL", "72 720 Td"]
        for j, line in enumerate(lines):
            if j:
                parts.append("T*")
            parts.append(f"({esc(line)}) Tj")
        parts.append("ET")
        stream = "\n".join(parts)
        body[content_ids[i]] = f"<< /Length {len(stream)} >>\nstream\n{stream}\nendstream"
    body[font_id] = "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>"
    out = bytearray(b"%PDF-1.4\n")
    offsets = {}
    for num in sorted(body):
        offsets[num] = len(out)
        out += f"{num} 0 obj\n{body[num]}\nendobj\n".encode("latin-1")
    xref_pos = len(out)
    count = len(body) + 1
    out += f"xref\n0 {count}\n".encode()
    out += b"0000000000 65535 f \n"
    for num in sorted(body):
        out += f"{offsets[num]:010d} 00000 n \n".encode()
    out += (f"trailer\n<< /Size {count} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF\n").encode()
    with open(path, "wb") as f:
        f.write(bytes(out))

def make_png(path, w=8, h=8, rgb=(90, 140, 60)):
    # Tiny handmade PNG for the add_picture demo.
    def chunk(tag, data):
        c = struct.pack('>I', len(data)) + tag + data
        return c + struct.pack('>I', zlib.crc32(tag + data) & 0xffffffff)
    ihdr = struct.pack('>IIBBBBB', w, h, 8, 2, 0, 0, 0)
    raw = b''.join(b'\x00' + bytes(rgb) * w for _ in range(h))
    png = (b'\x89PNG\r\n\x1a\n' + chunk(b'IHDR', ihdr)
           + chunk(b'IDAT', zlib.compress(raw)) + chunk(b'IEND', b''))
    with open(path, 'wb') as f:
        f.write(png)

PAGE0 = 'Harbour Rowing Club\nCommittee minutes, 12 March\nPresent: five members'

def lib_checks():
    from pypdf import PdfReader, PdfWriter
    import docx
    from docx.enum.text import WD_BREAK

    # --- sample files, exactly as built for the chapter ---
    make_pdf("minutes.pdf", [
        ["Harbour Rowing Club", "Committee minutes, 12 March", "Present: five members"],
        ["Item 1: boat shed roof", "Quotes to be gathered by April"],
        ["Item 2: regatta date", "Agreed for the last Saturday of June"],
    ])
    make_pdf("january.pdf", [
        ["Allotment News, January", "Frost cover reminders"],
        ["Seed catalogue highlights"],
    ])
    make_pdf("february.pdf", [
        ["Allotment News, February", "Pruning the fruit trees"],
    ])
    make_pdf("draft_mark.pdf", [["DRAFT COPY"]], font_size=48)

    print("=== section: Reading a PDF with pypdf ===")
    reader = PdfReader('minutes.pdf')
    check("len(reader.pages)", len(reader.pages), 3)
    check("pages[0].extract_text() repr", reader.pages[0].extract_text(), PAGE0)
    check("loop printing every page",
        run("""from pypdf import PdfReader
reader = PdfReader('minutes.pdf')
for page in reader.pages:
    print(page.extract_text())"""),
        "Harbour Rowing Club\nCommittee minutes, 12 March\nPresent: five members\n"
        "Item 1: boat shed roof\nQuotes to be gathered by April\n"
        "Item 2: regatta date\nAgreed for the last Saturday of June")

    print("=== section: New PDFs from old pages ===")
    run("""from pypdf import PdfReader, PdfWriter
writer = PdfWriter()
writer.append('january.pdf')
writer.append('february.pdf')
with open('winter.pdf', 'wb') as f:
    writer.write(f)""")
    check("merged page count", len(PdfReader('winter.pdf').pages), 3)
    run("""from pypdf import PdfReader, PdfWriter
writer = PdfWriter()
writer.append('january.pdf', pages=[0])
writer.append('february.pdf', pages=[0])
with open('front_pages.pdf', 'wb') as f:
    writer.write(f)""")
    check("selected-pages count", len(PdfReader('front_pages.pdf').pages), 2)
    check("split loop print lines",
        run("""from pypdf import PdfReader, PdfWriter
reader = PdfReader('winter.pdf')
for i, page in enumerate(reader.pages):
    writer = PdfWriter()
    writer.add_page(page)
    with open(f'winter_page_{i + 1}.pdf', 'wb') as f:
        writer.write(f)
    print(f'wrote winter_page_{i + 1}.pdf')"""),
        "wrote winter_page_1.pdf\nwrote winter_page_2.pdf\nwrote winter_page_3.pdf")
    for n in (1, 2, 3):
        check(f"winter_page_{n}.pdf has one page",
              len(PdfReader(f"winter_page_{n}.pdf").pages), 1)
    page = PdfReader('minutes.pdf').pages[0]
    check("rotation before", page.rotation, 0)
    check("chained rotate(90).rotation", page.rotate(90).rotation, 90)
    w = PdfWriter(); w.add_page(page)
    with open('minutes_rotated.pdf', 'wb') as f:
        w.write(f)
    check("prose claim: rotation survives write and reopen",
          PdfReader('minutes_rotated.pdf').pages[0].rotation, 90)
    page.rotate(90); page.rotate(90)
    check("prose claim: two further quarter-turns reach 270", page.rotation, 270)

    print("=== section: Watermarks and passwords ===")
    run("""from pypdf import PdfReader, PdfWriter
mark = PdfReader('draft_mark.pdf').pages[0]
writer = PdfWriter()
for page in PdfReader('minutes.pdf').pages:
    page.merge_page(mark)
    writer.add_page(page)
with open('minutes_draft.pdf', 'wb') as f:
    writer.write(f)""")
    check("stamped page carries both layers' text",
          PdfReader('minutes_draft.pdf').pages[0].extract_text(),
          PAGE0 + '\nDRAFT COPY')
    run("""from pypdf import PdfReader, PdfWriter
writer = PdfWriter()
writer.append('minutes.pdf')
writer.encrypt('swordfish')
with open('minutes_locked.pdf', 'wb') as f:
    writer.write(f)""")
    reader = PdfReader('minutes_locked.pdf')
    check("is_encrypted", reader.is_encrypted, True)
    try:
        reader.pages[0].extract_text()
        got = "no error raised"
    except Exception as ex:
        got = f"{type(ex).__name__}: {ex}"
    check("error before decrypt (single final line shown in chapter)",
          got, "FileNotDecryptedError: File has not been decrypted")
    ret = reader.decrypt('swordfish')
    check("decrypt return repr", repr(ret), "<PasswordType.OWNER_PASSWORD: 2>")
    check("prose claim: correct password is truthy", bool(ret), True)
    check("extract after decrypt", reader.pages[0].extract_text(), PAGE0)
    bad = PdfReader('minutes_locked.pdf').decrypt('guess')
    check("wrong password repr", repr(bad), "<PasswordType.NOT_DECRYPTED: 0>")
    check("prose claim: wrong password is falsy", bool(bad), False)
    check("note claim: file on disk stays encrypted",
          PdfReader('minutes_locked.pdf').is_encrypted, True)

    print("=== sections: Word documents (read and write) ===")
    run("""import docx
doc = docx.Document()
doc.add_heading('The Plot Holder', 0)
doc.add_heading('Compost rota', 1)
doc.add_paragraph('Turning duty passes to bed seven this month.')
p = doc.add_paragraph('Deliveries arrive ')
p.add_run('Thursday').bold = True
p.add_run(' at the north gate.')
doc.add_heading('Seed swap', 1)
doc.add_paragraph('Bring labelled envelopes to the shed.')
doc.save('newsletter.docx')""")
    d = docx.Document('newsletter.docx')
    check("len(doc.paragraphs)", len(d.paragraphs), 6)
    check("paragraphs[0].text", d.paragraphs[0].text, 'The Plot Holder')
    check("paragraphs[3].text", d.paragraphs[3].text,
          'Deliveries arrive Thursday at the north gate.')
    para = d.paragraphs[3]
    check("len(para.runs)", len(para.runs), 3)
    check("runs[1].text", para.runs[1].text, 'Thursday')
    check("runs[1].bold", para.runs[1].bold, True)
    check("runs[0].bold prints None", para.runs[0].bold, None)
    check("runs[2].bold is None too (prose: plain neighbours)", para.runs[2].bold, None)
    check("style of paragraphs[0]", d.paragraphs[0].style.name, 'Title')
    check("style of paragraphs[1]", d.paragraphs[1].style.name, 'Heading 1')
    check("style of paragraphs[2]", d.paragraphs[2].style.name, 'Normal')
    check("style of paragraphs[4]", d.paragraphs[4].style.name, 'Heading 1')

    # runs_demo.docx: italics switched on mid-word, as described in section 5.
    doc2 = docx.Document()
    q = doc2.add_paragraph('')
    q.add_run('Wat')
    q.add_run('er butts').italic = True
    q.add_run(' refill on Fridays.')
    doc2.save('runs_demo.docx')
    para = docx.Document('runs_demo.docx').paragraphs[0]
    check("runs_demo para.text", para.text, 'Water butts refill on Fridays.')
    check("runs_demo run pieces", [r.text for r in para.runs],
          ['Wat', 'er butts', ' refill on Fridays.'])
    check("runs_demo italics tri-state", [r.italic for r in para.runs],
          [None, True, None])

    check("full_text function output",
        run("""import docx

def full_text(path):
    doc = docx.Document(path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return '\\n'.join(text)

print(full_text('newsletter.docx'))"""),
        "The Plot Holder\nCompost rota\nTurning duty passes to bed seven this month.\n"
        "Deliveries arrive Thursday at the north gate.\nSeed swap\nBring labelled envelopes to the shed.")
    check("Heading 1 filter output",
        run("""import docx
doc = docx.Document('newsletter.docx')
for para in doc.paragraphs:
    if para.style.name == 'Heading 1':
        print(para.text)"""),
        "Compost rota\nSeed swap")

    check("fresh Document has 0 paragraphs", len(docx.Document().paragraphs), 0)
    run("""import docx
from docx.enum.text import WD_BREAK
doc = docx.Document()
doc.add_paragraph('End of part one.')
doc.paragraphs[0].runs[0].add_break(WD_BREAK.PAGE)
doc.add_paragraph('Part two begins.')
doc.save('two_parts.docx')""")
    d4 = docx.Document('two_parts.docx')
    check("page break leaves paragraph count at two", len(d4.paragraphs), 2)
    check("two_parts texts", [p.text for p in d4.paragraphs],
          ['End of part one.', 'Part two begins.'])

    make_png('sprout.png')
    run("""import docx
doc = docx.Document()
doc.add_paragraph('The first seedling of spring:')
doc.add_picture('sprout.png', width=docx.shared.Inches(1))
doc.save('picture_demo.docx')""")
    d5 = docx.Document('picture_demo.docx')
    check("len(doc.inline_shapes)", len(d5.inline_shapes), 1)
    check("shape width echoes as 914400", repr(d5.inline_shapes[0].width), '914400')
    check("Inches(1) echoes as 914400", repr(docx.shared.Inches(1)), '914400')
    check("prose claim: height follows proportions (square image)",
          int(d5.inline_shapes[0].height), 914400)

    doc7 = docx.Document()
    doc7.add_heading('Level two', 2)
    doc7.save('h2_demo.docx')
    check("q10 claim: level 2 gives Heading 2",
          docx.Document('h2_demo.docx').paragraphs[0].style.name, 'Heading 2')

if not HAVE_LIBS:
    print(f"skip: pypdf/python-docx not installed ({GUARDED_CHECKS} checks skipped)")
else:
    before = total
    old = os.getcwd()
    with tempfile.TemporaryDirectory() as tmp:
        os.chdir(tmp)
        try:
            lib_checks()
        finally:
            os.chdir(old)
    ran = total - before
    check("guarded section ran the advertised check count", ran, GUARDED_CHECKS)

print()
print("CH17 VERIFY: ALL PASS" if fails == 0 else f"CH17 VERIFY: {fails} FAILURE(S)")
sys.exit(0 if fails == 0 else 1)
